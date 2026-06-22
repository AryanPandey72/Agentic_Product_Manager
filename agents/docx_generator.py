from docx import Document
from docx.shared import Inches


def add_bullet_section(doc, title, items):
    doc.add_heading(title, level=2)

    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def generate_product_blueprint(
    requirements,
    strategy,
    architecture,
    output_path
):
    doc = Document()

    # ==================================================
    # COVER PAGE
    # ==================================================

    doc.add_heading(
        "Product Blueprint",
        level=0
    )
    # ==================================================
    # EXECUTIVE SUMMARY
    # ==================================================

    doc.add_heading(
        "Executive Summary",
        level=1
    )

    target_users = ", ".join(
        requirements.target_users[:3]
    )

    business_goals = ", ".join(
        requirements.business_goals[:3]
    )

    executive_summary = (
        f"This document outlines a product designed to address "
        f"{requirements.problem_statement.lower()} "
        f"The solution is intended for {target_users} and aims to achieve "
        f"business objectives such as {business_goals}. "
        f"The proposed product strategy focuses on "
        f"{strategy.unique_value_proposition.lower()}"
    )

    doc.add_paragraph(
        executive_summary
    )
    # ==================================================
    # REQUIREMENTS ANALYSIS
    # ==================================================

    doc.add_heading(
        "Requirements Analysis",
        level=1
    )

    doc.add_heading(
        "Problem Statement",
        level=2
    )

    doc.add_paragraph(
        requirements.problem_statement
    )

    add_bullet_section(
        doc,
        "Target Users",
        requirements.target_users
    )

    add_bullet_section(
        doc,
        "Business Goals",
        requirements.business_goals
    )

    add_bullet_section(
        doc,
        "Constraints",
        requirements.constraints
    )

    add_bullet_section(
        doc,
        "Success Metrics",
        requirements.success_metrics
    )

    # ==================================================
    # PRODUCT STRATEGY
    # ==================================================

    doc.add_heading(
        "Product Strategy",
        level=1
    )

    doc.add_heading(
        "Market Positioning",
        level=2
    )

    doc.add_paragraph(
        strategy.market_positioning
    )

    doc.add_heading(
        "Unique Value Proposition",
        level=2
    )

    doc.add_paragraph(
        strategy.unique_value_proposition
    )

    add_bullet_section(
        doc,
        "Monetization Strategy",
        strategy.monetization_strategy
    )

    add_bullet_section(
        doc,
        "Acquisition Channels",
        strategy.acquisition_channels
    )

    add_bullet_section(
        doc,
        "MVP Scope",
        strategy.mvp_scope
    )

    add_bullet_section(
        doc,
        "Go-To-Market Phases",
        strategy.gtm_phases
    )

    doc.add_heading(
        "Risks & Mitigations",
        level=2
    )

    for item in strategy.risks_and_mitigations:

        doc.add_paragraph(
            f"Risk: {item.risk}",
            style="List Bullet"
        )

        doc.add_paragraph(
            f"Mitigation: {item.mitigation}",
            style="List Bullet 2"
        )

    # ==================================================
    # TECHNICAL ARCHITECTURE
    # ==================================================

    doc.add_heading(
        "Technical Architecture",
        level=1
    )

    add_bullet_section(
        doc,
        "Technology Stack",
        architecture.tech_stack
    )

    doc.add_heading(
        "Infrastructure",
        level=2
    )

    doc.add_paragraph(
        f"Hosting: {architecture.infrastructure.hosting}"
    )

    doc.add_paragraph(
        f"Caching: {architecture.infrastructure.caching_strategy}"
    )

    doc.add_paragraph(
        f"Background Workers: {architecture.infrastructure.background_workers}"
    )

    doc.add_heading(
        "Security",
        level=2
    )

    doc.add_paragraph(
        f"Authentication: {architecture.security.authentication_method}"
    )

    doc.add_paragraph(
        f"Roles: {', '.join(architecture.security.authorization_roles)}"
    )

    doc.add_paragraph(
        f"Data Privacy: {architecture.security.data_privacy}"
    )

    add_bullet_section(
        doc,
        "Third Party Integrations",
        architecture.third_party_integrations
    )

    # ==================================================
    # DATABASE DESIGN
    # ==================================================

    doc.add_heading(
        "Database Design",
        level=1
    )

    for table in architecture.database_schema:

        doc.add_heading(
            table.table_name,
            level=2
        )

        database_table = doc.add_table(
            rows=1,
            cols=2
        )

        database_table.style = "Table Grid"

        header = database_table.rows[0].cells

        header[0].text = "Column"
        header[1].text = "Data Type"

        for column in table.columns:

            row = database_table.add_row().cells

            row[0].text = column.name
            row[1].text = column.data_type

        if table.relationships:

            doc.add_paragraph(
                "Relationships:",
                style=None
            )

            for relation in table.relationships:

                doc.add_paragraph(
                    relation,
                    style="List Bullet"
                )

    # ==================================================
    # API DESIGN
    # ==================================================

    doc.add_heading(
        "API Design",
        level=1
    )

    api_table = doc.add_table(
        rows=1,
        cols=3
    )

    api_table.style = "Table Grid"

    header = api_table.rows[0].cells

    header[0].text = "Method"
    header[1].text = "Endpoint"
    header[2].text = "Purpose"

    for api in architecture.api_specifications:

        row = api_table.add_row().cells

        row[0].text = api.method
        row[1].text = api.endpoint
        row[2].text = api.purpose

    # ==================================================
    # FEATURE TRACEABILITY MATRIX
    # ==================================================

    doc.add_heading(
        "Feature Traceability Matrix",
        level=1
    )

    matrix = doc.add_table(
        rows=1,
        cols=3
    )

    matrix.style = "Table Grid"

    header = matrix.rows[0].cells

    header[0].text = "Feature"
    header[1].text = "Database Tables"
    header[2].text = "API Endpoints"

    for feature in architecture.feature_mappings:

        row = matrix.add_row().cells

        row[0].text = feature.feature_name

        row[1].text = "\n".join(
            feature.database_tables
        )

        row[2].text = "\n".join(
            feature.api_endpoints
        )

    # ==================================================
    # SAVE
    # ==================================================

    doc.save(output_path)

    return output_path