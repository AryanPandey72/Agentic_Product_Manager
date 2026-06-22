from docx import Document


def generate_sprint_plan_docx(
    sprint_plan,
    output_path
):

    doc = Document()

    doc.add_heading(
        "Sprint Plan",
        level=0
    )

    for sprint in sprint_plan.sprints:

        doc.add_heading(
            f"Sprint {sprint.sprint_number}",
            level=1
        )

        doc.add_paragraph(
            f"Goal: {sprint.sprint_goal}"
        )

        for epic in sprint.epics:

            doc.add_heading(
                epic.title,
                level=2
            )

            doc.add_paragraph(
                epic.description
            )

            for story in epic.stories:

                doc.add_heading(
                    story.title,
                    level=3
                )

                doc.add_paragraph(
                    story.description
                )

                doc.add_paragraph(
                    "Acceptance Criteria:"
                )

                for criterion in story.acceptance_criteria:

                    doc.add_paragraph(
                        criterion,
                        style="List Bullet"
                    )

        doc.add_page_break()

    doc.save(
        output_path
    )